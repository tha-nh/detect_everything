from __future__ import annotations

import re
from dataclasses import dataclass
from html import escape
from collections import Counter
from pathlib import Path


TEXT_EXTENSIONS = {".txt", ".docx", ".pdf"}


@dataclass(frozen=True)
class LabelDefinition:
    name: str
    description: str
    tokens: set[str]


@dataclass(frozen=True)
class LabeledSegment:
    text: str
    label: str
    score: float


class TextFileDetector:
    """Extract text from document files and classify text segments into labels."""

    def __init__(self, keywords: list[str]) -> None:
        cleaned = [keyword.strip() for keyword in keywords if keyword.strip()]
        if not cleaned:
            raise ValueError("Keyword list cannot be empty.")

        self.labels = [self._parse_label_definition(item) for item in cleaned]
        self.semantic_model = self._load_semantic_model()

    def detect_file(self, input_path: str, output_path: str) -> dict[str, int]:
        source = Path(input_path)
        if not source.is_file():
            raise FileNotFoundError(f"File not found: {source}")

        text = self.extract_text(source)
        labeled_segments = self.classify_text(text)
        counts: Counter[str] = Counter(segment.label for segment in labeled_segments)

        destination = Path(output_path)
        destination.parent.mkdir(parents=True, exist_ok=True)
        destination.write_text(
            self.build_report(source, text, labeled_segments, counts),
            encoding="utf-8",
        )

        return dict(counts)

    def extract_text(self, source: Path) -> str:
        extension = source.suffix.lower()

        if extension == ".txt":
            return self._extract_txt(source)
        if extension == ".docx":
            return self._extract_docx(source)
        if extension == ".pdf":
            return self._extract_pdf(source)

        raise ValueError(f"Unsupported text file: {source}")

    def classify_text(self, text: str) -> list[LabeledSegment]:
        segments = self._split_segments(text)
        if self.semantic_model is not None:
            return self._classify_segments_semantically(segments)
        return [self._classify_segment(segment) for segment in segments]

    def _classify_segments_semantically(self, segments: list[str]) -> list[LabeledSegment]:
        if not segments:
            return []

        label_texts = [
            f"{definition.name}. {definition.description}"
            for definition in self.labels
        ]
        label_embeddings = self.semantic_model.encode(
            label_texts,
            convert_to_tensor=True,
            normalize_embeddings=True,
        )
        segment_embeddings = self.semantic_model.encode(
            segments,
            convert_to_tensor=True,
            normalize_embeddings=True,
        )

        try:
            from sentence_transformers import util
        except ImportError:
            return [self._classify_segment(segment) for segment in segments]

        scores = util.cos_sim(segment_embeddings, label_embeddings)
        labeled_segments: list[LabeledSegment] = []
        for segment, row in zip(segments, scores):
            best_index = int(row.argmax().item())
            best_score = float(row[best_index].item())
            label = self.labels[best_index].name if best_score >= 0.22 else "unlabeled"
            labeled_segments.append(LabeledSegment(segment, label, best_score))

        return labeled_segments

    def _classify_segment(self, segment: str) -> LabeledSegment:
        segment_tokens = self._tokenize(segment)
        if not segment_tokens:
            return LabeledSegment(segment, "unlabeled", 0.0)

        best_label = "unlabeled"
        best_score = 0.0

        for definition in self.labels:
            overlap = segment_tokens & definition.tokens
            token_score = len(overlap) / max(1, len(definition.tokens))
            phrase_score = 0.25 if definition.name.casefold() in segment.casefold() else 0.0
            score = min(1.0, token_score + phrase_score)
            if score > best_score:
                best_score = score
                best_label = definition.name

        if best_score <= 0:
            best_label = "unlabeled"

        return LabeledSegment(segment, best_label, best_score)

    def build_report(
        self,
        source: Path,
        text: str,
        labeled_segments: list[LabeledSegment],
        counts: Counter[str],
    ) -> str:
        label_rows = "\n".join(
            "<tr>"
            f"<td>{escape(definition.name)}</td>"
            f"<td>{escape(definition.description)}</td>"
            f"<td>{counts.get(definition.name, 0)}</td>"
            "</tr>"
            for definition in self.labels
        )
        extracted_text = text.strip()
        if not extracted_text:
            extracted_text = escape(
                "No extractable text was found. If this is a scanned PDF, "
                "OCR support is needed."
            )
        segment_blocks = "\n".join(
            '<article class="segment">'
            f'<div class="segment-label">{escape(segment.label)}'
            f'<span>{segment.score:.0%} match</span></div>'
            f"<p>{escape(segment.text)}</p>"
            "</article>"
            for segment in labeled_segments
        ) or f'<article class="segment"><p>{extracted_text}</p></article>'

        return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Detected Text - {escape(source.name)}</title>
  <style>
    :root {{
      color-scheme: light;
      --bg: #f6f7f9;
      --surface: #ffffff;
      --text: #172033;
      --muted: #667085;
      --border: #d7dde8;
      --label: #0f766e;
      --label-soft: #ccfbf1;
    }}
    body {{
      margin: 0;
      background: var(--bg);
      color: var(--text);
      font-family: "Segoe UI", Arial, sans-serif;
      line-height: 1.55;
    }}
    main {{
      max-width: 980px;
      margin: 0 auto;
      padding: 28px 20px 40px;
    }}
    h1 {{
      margin: 0 0 6px;
      font-size: 26px;
    }}
    .meta {{
      margin: 0 0 22px;
      color: var(--muted);
      font-size: 14px;
    }}
    section {{
      background: var(--surface);
      border: 1px solid var(--border);
      border-radius: 8px;
      padding: 18px;
      margin-top: 16px;
    }}
    h2 {{
      margin: 0 0 12px;
      font-size: 17px;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      font-size: 14px;
    }}
    th, td {{
      border-bottom: 1px solid var(--border);
      padding: 9px 10px;
      text-align: left;
    }}
    th {{
      color: var(--muted);
      font-weight: 700;
      background: #f8fafc;
    }}
    pre {{
      margin: 0;
      white-space: pre-wrap;
      word-wrap: break-word;
      font: 15px/1.7 Consolas, "Segoe UI", monospace;
    }}
    .segment {{
      border: 1px solid var(--border);
      border-radius: 8px;
      padding: 14px;
      margin: 12px 0;
      background: #fbfcfe;
    }}
    .segment p {{
      margin: 10px 0 0;
      white-space: pre-wrap;
    }}
    .segment-label {{
      display: inline-flex;
      align-items: center;
      gap: 8px;
      border-radius: 999px;
      background: var(--label-soft);
      color: var(--label);
      padding: 4px 10px;
      font-weight: 800;
      font-size: 13px;
    }}
    .segment-label span {{
      color: var(--muted);
      font-weight: 700;
      font-size: 12px;
    }}
  </style>
</head>
<body>
  <main>
    <h1>Detected Text Labels</h1>
    <p class="meta">Source: {escape(str(source))} | Type: {escape(source.suffix.lower())}</p>

    <section>
      <h2>Labels</h2>
      <table>
        <thead>
          <tr><th>Label</th><th>Description</th><th>Segments</th></tr>
        </thead>
        <tbody>
          {label_rows}
        </tbody>
      </table>
    </section>

    <section>
      <h2>Labeled Segments</h2>
      {segment_blocks}
    </section>
  </main>
</body>
</html>
"""

    @classmethod
    def _parse_label_definition(cls, value: str) -> LabelDefinition:
        if "=" in value:
            name, description = value.split("=", 1)
        elif ":" in value:
            name, description = value.split(":", 1)
        else:
            name = value
            description = value

        clean_name = name.strip()
        clean_description = description.strip() or clean_name
        tokens = cls._tokenize(f"{clean_name} {clean_description}")
        return LabelDefinition(clean_name, clean_description, tokens)

    @staticmethod
    def _split_segments(text: str) -> list[str]:
        segments = [
            segment.strip()
            for segment in re.split(r"\n\s*\n+", text)
            if segment.strip()
        ]
        if segments:
            return segments

        return [line.strip() for line in text.splitlines() if line.strip()]

    @staticmethod
    def _tokenize(value: str) -> set[str]:
        stopwords = {
            "a",
            "an",
            "and",
            "are",
            "as",
            "cua",
            "của",
            "for",
            "in",
            "is",
            "la",
            "là",
            "of",
            "or",
            "the",
            "to",
            "va",
            "và",
            "voi",
            "với",
        }
        tokens = {
            token.casefold()
            for token in re.findall(r"[\wÀ-ỹ]+", value, flags=re.UNICODE)
            if len(token) > 1
        }
        return tokens - stopwords

    @staticmethod
    def _safe_label(value: str) -> str:
        label = re.sub(r"[^a-zA-Z0-9_-]+", "_", value.strip().lower())
        return label.strip("_") or "text"

    @staticmethod
    def _load_semantic_model():
        try:
            from sentence_transformers import SentenceTransformer
        except ImportError:
            return None

        return SentenceTransformer("sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")

    @staticmethod
    def _extract_txt(source: Path) -> str:
        last_error: UnicodeDecodeError | None = None
        for encoding in ("utf-8-sig", "utf-8", "cp1258", "cp1252"):
            try:
                return source.read_text(encoding=encoding)
            except UnicodeDecodeError as exc:
                last_error = exc

        if last_error is not None:
            raise last_error
        return ""

    @staticmethod
    def _extract_docx(source: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "Missing dependency python-docx. Run: pip install -r requirements.txt"
            ) from exc

        document = Document(str(source))
        chunks: list[str] = []
        chunks.extend(paragraph.text for paragraph in document.paragraphs)

        for table in document.tables:
            for row in table.rows:
                chunks.append("\t".join(cell.text for cell in row.cells))

        return "\n".join(chunk for chunk in chunks if chunk.strip())

    @staticmethod
    def _extract_pdf(source: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError(
                "Missing dependency pypdf. Run: pip install -r requirements.txt"
            ) from exc

        reader = PdfReader(str(source))
        pages: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"--- Page {index} ---\n{page_text.strip()}")

        return "\n\n".join(pages)
